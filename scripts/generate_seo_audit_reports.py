from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


DATE = "2026-07-19"
OUT = Path(__file__).resolve().parents[1] / "docs" / "seo-audits" / DATE

NAVY = "1B2A4A"
BLUE = "2563EB"
GREEN = "16A34A"
AMBER = "D97706"
RED = "DC2626"
LIGHT_BLUE = "EFF6FF"
LIGHT_GREEN = "F0FDF4"
LIGHT_GRAY = "F8F9FA"
MID_GRAY = "E2E8F0"
DARK = "1E293B"


SITES = [
    {
        "slug": "sanfordcleaning-com",
        "name": "Sanford Cleaning",
        "domain": "sanfordcleaning.com",
        "scores": (7, 8, 8),
        "pages": "37 pages: homepage, About, FAQ, Guides index, 11 service pages, booking/custom quote/careers, and 15 guide articles; robots.txt, sitemap.xml, and llms.txt also inspected.",
        "summary": "Sanford has the strongest AI-search foundation in the portfolio: rich local entity schema, FAQPage coverage, HowTo markup, Speakable targets, and a detailed llms.txt. The main risk is factual inconsistency: the site gives conflicting pricing floors ($80 versus $140–$350+) and conflicting experience claims (10+ years versus 2 years under the current company). Local identity is incomplete because no street address is published, and four guide pages are thin and carry only breadcrumb schema.",
        "findings": [
            ("SEO", "Title tags", "34 of 37 audited titles exceed 60 characters; most key pages are likely to truncate.", "Needs Attention"),
            ("SEO", "Technical hygiene", "HTTPS, HSTS/CSP, self-canonicals, index/follow, viewport, OG/Twitter and descriptive image alts are present.", "Good"),
            ("SEO", "Content depth", "Major service pages contain roughly 850–1,400 words and are well interlinked.", "Good"),
            ("SEO", "Fact consistency", "Homepage/FAQ advertise $80 starts while pricing guides state $140–$350+ and $170–$240 packages.", "Critical"),
            ("SEO", "Thin content", "Four comparison/deals guides contain only about 343–400 words and BreadcrumbList schema.", "Needs Attention"),
            ("GEO", "AI discoverability", "llms.txt provides brand facts, pricing, CTAs and entities; three guides expose Speakable targets.", "Good"),
            ("GEO", "E-E-A-T", "Topaz West LLC is named, but authors are generic and no named team biographies or review evidence were found.", "Needs Attention"),
            ("GEO", "Entity consistency", "10+ years, 2 years, and 5–7 years appear across pages, weakening machine trust.", "Critical"),
            ("AEO", "Structured answers", "FAQPage appears on roughly 27 pages; HowTo exists on Airbnb, move-out and booking content.", "Good"),
            ("AEO", "Voice/local", "Phone, email, hours and service areas are present, but streetAddress is absent.", "Needs Attention"),
        ],
        "priorities": [
            ("Critical", "Unify all pricing statements and schema answers around one source of truth.", "SEO/GEO/AEO", "Medium", "Very high"),
            ("Critical", "Reconcile years-in-business claims with a precise company/history statement.", "GEO", "Low", "High"),
            ("High", "Publish complete NAP or explicitly document a service-area/mobile-business address policy.", "Local/AEO", "Low–Medium", "High"),
            ("High", "Expand four thin guides to 800–1,200 words with tables, FAQs and Article schema.", "SEO/AEO", "Medium", "High"),
            ("Medium", "Shorten titles to roughly 50–60 characters.", "SEO", "Low", "Medium"),
            ("Medium", "Add named expert authors and verified review evidence.", "GEO", "Medium", "High"),
        ],
        "strengths": ["Production-grade llms.txt", "FAQ/HowTo/Speakable coverage", "First-party Sanford pricing and humidity content", "Strong security and crawlability"],
    },
    {
        "slug": "deltonacleaning-com",
        "name": "Deltona Cleaning",
        "domain": "deltonacleaning.com",
        "scores": (7, 6, 8),
        "pages": "15 pages: homepage, About, Services, Contact, Blog index, 9 individual services, and the move-out inspection guide; robots.txt and sitemap.xml also inspected.",
        "summary": "Deltona has a clean local service architecture, consistent geo targeting, Service and FAQ schema, and one genuinely citable move-out inspection guide. AEO is the strongest dimension. GEO is held back by a thin About page, no named people or credentials, no street address, no sameAs entity links, and no external citations. The homepage and Services hub also share a title, and the content program currently consists of a single article.",
        "findings": [
            ("SEO", "Titles", "Homepage and /services share the same title; Contact repeats the brand suffix; the guide title is 103 characters.", "Needs Attention"),
            ("SEO", "Meta descriptions", "Present sitewide, but several key pages exceed 160 characters.", "Needs Attention"),
            ("SEO", "Service architecture", "Nine descriptive service URLs use singular H1s, clean paths and related-service links.", "Good"),
            ("SEO", "Editorial depth", "Only one blog post was discovered; About and Contact are thin.", "Needs Attention"),
            ("SEO", "Structured data", "Organization, ProfessionalService, Service, FAQPage, BreadcrumbList and TechArticle are implemented.", "Good"),
            ("GEO", "People and expertise", "No named owner, manager, cleaner or credentialed Person entity appears.", "Missing"),
            ("GEO", "Original data", "The move-out guide publishes fail-rate figures such as oven 41% and grout 37%.", "Good"),
            ("GEO", "Entity graph", "No sameAs links to GBP or social profiles were found.", "Missing"),
            ("AEO", "Answer formatting", "Question headings, FAQ answers, lists and a comparison table are well suited to extraction.", "Good"),
            ("AEO", "HowTo/Speakable", "Visible process steps are not semantic ordered lists and have no HowTo or Speakable schema.", "Missing"),
        ],
        "priorities": [
            ("High", "Add named leadership/team information, credentials and Person schema.", "GEO", "Medium", "High"),
            ("High", "Complete NAP and link a verified Google Business Profile through sameAs.", "Local/AEO", "Low–Medium", "High"),
            ("High", "Publish 4–6 additional Deltona-specific, data-backed guides.", "SEO/GEO", "Medium–High", "High"),
            ("High", "Remove duplicate titles and shorten the 103-character guide title.", "SEO", "Low", "Medium"),
            ("Medium", "Convert visible service processes to semantic ol + HowTo JSON-LD.", "AEO", "Low–Medium", "Medium"),
            ("Medium", "Cite authoritative Florida/local sources in informational guides.", "GEO", "Medium", "Medium"),
        ],
        "strengths": ["Nine-page service architecture", "Broad FAQPage implementation", "Original move-out inspection data", "Strong local entity and OfferCatalog schema"],
    },
    {
        "slug": "hainescitycleaning-com",
        "name": "Haines City Cleaning",
        "domain": "hainescitycleaning.com",
        "scores": (7, 5, 7),
        "pages": "All 16 sitemap URLs: homepage, Pricing, Blog index, 7 service pages, 3 area pages and 3 blog posts; robots.txt and sitemap.xml also inspected.",
        "summary": "Haines City has solid traditional SEO, transparent prices, useful local service pages, FAQPage markup, and fact-dense Airbnb and move-out articles. The largest risk is invalid sitewide entity data: LocalBusiness sameAs values still contain literal PLACEHOLDER URLs. No About or Team page exists, no street address is published, and the three area pages contain only about 200 words each. These weaknesses explain the lower GEO score despite a strong answer-ready content foundation.",
        "findings": [
            ("SEO", "On-page basics", "All crawled pages have a single H1, self-canonical, descriptive path, indexability and OG/Twitter metadata.", "Good"),
            ("SEO", "Title/meta length", "Multiple titles are 78–101 characters and several meta descriptions exceed 160 characters.", "Needs Attention"),
            ("SEO", "Area content", "Davenport, Winter Haven and Lakeland area pages contain only about 196–202 words.", "Needs Attention"),
            ("SEO", "Pricing", "Visible starting rates and add-ons create useful commercial and answer-engine signals.", "Good"),
            ("SEO", "Structured data", "LocalBusiness/HouseCleaner, ProfessionalService, FAQPage, BreadcrumbList and BlogPosting are present.", "Good"),
            ("GEO", "sameAs integrity", "Schema contains PLACEHOLDER_GOOGLE_BUSINESS_PROFILE and placeholder social/BBB URLs.", "Critical"),
            ("GEO", "E-E-A-T", "About, Team and Our Story routes return 404; blog authors are Organization-only.", "Missing"),
            ("GEO", "Original facts", "Airbnb content publishes 90–120 minute two-person turnover SLAs and phase timings.", "Good"),
            ("AEO", "FAQ coverage", "Core services and two blog posts include real FAQPage questions and answers.", "Good"),
            ("AEO", "HowTo/Speakable", "Checklist and SLA content has no HowTo or Speakable markup.", "Missing"),
        ],
        "priorities": [
            ("Critical", "Replace or remove every placeholder sameAs/GBP/social/BBB URL.", "SEO/GEO", "Low", "Very high"),
            ("Critical", "Create a real About page with named people, experience and credentials.", "GEO", "Medium", "Very high"),
            ("High", "Complete NAP with a real address or explicit service-area policy.", "Local/AEO", "Low", "High"),
            ("High", "Expand all area pages with unique neighborhoods, housing context and FAQs.", "SEO/GEO", "Medium", "High"),
            ("High", "Expand the thin cost guide and move-in service page.", "SEO/AEO", "Medium", "High"),
            ("Medium", "Add HowTo and table markup to checklist/pricing content.", "AEO", "Low–Medium", "Medium"),
        ],
        "strengths": ["Transparent pricing", "Strong FAQ schema footprint", "Local Airbnb and move-out facts", "Complete 16-URL crawl and clean internal structure"],
    },
    {
        "slug": "celebrationcleaning-com",
        "name": "Celebration Cleaning",
        "domain": "celebrationcleaning.com",
        "scores": (5, 6, 5),
        "pages": "145 sitemap URLs: 5 core/content pages, 14 city hubs and 126 city×service pages (14 cities × 9 services); both apex/www robots and sitemap variants were inspected.",
        "summary": "Celebration has the portfolio's largest programmatic footprint, but also its highest SEO risk. Live traffic redirects www to the apex host while every canonical and sitemap URL declares www, sending contradictory indexing signals. All 126 city×service pages are thin near-duplicates of roughly 267 words, share only nine body templates, and visibly state that the lead form is a demonstration. The Florida Airbnb turnover guide is the standout asset and should become the model for rebuilding local pages.",
        "findings": [
            ("SEO", "Canonical host", "www redirects to apex, while canonical tags, sitemap URLs and schema identify www.", "Critical"),
            ("SEO", "Programmatic quality", "126 city×service pages reduce to nine templates with city-name substitution.", "Critical"),
            ("SEO", "Production readiness", "All 126 local service pages display '*This is a demonstration form'.", "Critical"),
            ("SEO", "Core metadata", "Unique titles, descriptions, singular H1s, viewport and OG metadata exist across the crawl.", "Good"),
            ("SEO", "Schema", "Organization and ProfessionalService are present, but BreadcrumbList is absent across 145 pages.", "Needs Attention"),
            ("GEO", "E-E-A-T", "About has no named people; no street address/email; several sameAs destinations are broken or generic.", "Needs Attention"),
            ("GEO", "Original evidence", "The Airbnb guide publishes median turnover minutes, success percentages and linen-par data.", "Good"),
            ("AEO", "Sitewide answers", "Only the Airbnb guide has FAQPage; local service pages provide no search-question answers.", "Missing"),
            ("AEO", "Advanced markup", "No HowTo or Speakable markup exists across the crawl.", "Missing"),
            ("AEO", "Local voice", "Phone and hours exist, but no full NAP/verified GBP alignment was confirmed.", "Needs Attention"),
        ],
        "priorities": [
            ("Critical", "Choose apex or www and align redirects, canonicals, sitemap, schema and OG URLs.", "SEO", "Medium", "Very high"),
            ("Critical", "Replace demonstration forms and deepen, consolidate or noindex the 126 thin pages.", "SEO/Trust", "High", "Very high"),
            ("High", "Complete NAP and replace broken sameAs destinations with verified profiles.", "GEO/AEO", "Low–Medium", "High"),
            ("High", "Expand About with named leadership, founding facts and training/insurance evidence.", "GEO", "Medium", "High"),
            ("High", "Add FAQ blocks/schema to city hubs and high-value service pages.", "AEO", "Medium", "High"),
            ("Medium", "Add BreadcrumbList, a true logo asset and a valid search action or remove it.", "SEO", "Low–Medium", "Medium"),
        ],
        "strengths": ["Complete multi-city crawlable architecture", "Clear package pricing", "Ambitious Organization/OfferCatalog graph", "Excellent Florida Airbnb operations guide"],
    },
    {
        "slug": "cleaningwinterhaven-com",
        "name": "Cleaning Winter Haven",
        "domain": "cleaningwinterhaven.com",
        "scores": (7, 4, 7),
        "pages": "9 sitemap URLs: homepage, 6 service pages, Pricing and Service Areas; robots.txt and sitemap.xml inspected. About, Team, Contact, FAQ and Blog routes were probed and returned 404.",
        "summary": "Cleaning Winter Haven has good service-page SEO, transparent pricing, FAQ schema, and strong neighborhood coverage. GEO is the weak point because robots.txt blocks GPTBot, Google-Extended, ClaudeBot and other major AI crawlers. There is no About/Team page, streetAddress is empty, and no sameAs graph exists. A separate technical issue is that HTTP serves content with status 200 instead of redirecting to HTTPS, creating duplicate-protocol risk.",
        "findings": [
            ("SEO", "Protocol canonicalization", "HTTP serves a 200 response instead of a permanent redirect to HTTPS; HSTS was not observed.", "Critical"),
            ("SEO", "Metadata", "Service titles/metas are mostly strong; homepage title is 70 characters and description is about 202.", "Needs Attention"),
            ("SEO", "OG sharing", "Every page reuses generic OG text and a 512px app icon rather than a 1200×630 page visual.", "Needs Attention"),
            ("SEO", "Content", "Service pages are useful; Pricing and Service Areas are thin at roughly 309 and 301 words.", "Needs Attention"),
            ("SEO", "Schema", "LocalBusiness, WebSite, Service, FAQPage and BreadcrumbList are implemented.", "Good"),
            ("GEO", "AI crawler access", "robots.txt disallows GPTBot, Google-Extended, ClaudeBot, Applebot-Extended and others.", "Critical"),
            ("GEO", "E-E-A-T", "No About/Team/people, credentials, sameAs links or street address were found.", "Missing"),
            ("GEO", "Citable facts", "Pricing bands, add-on percentages and neighborhood coverage are machine-summarizable.", "Good"),
            ("AEO", "FAQ coverage", "Homepage and all six service pages provide unique, concise FAQPage answers.", "Good"),
            ("AEO", "Advanced formats", "The three-phase construction workflow has no HowTo; Speakable is absent.", "Missing"),
        ],
        "priorities": [
            ("Critical", "Decide AI-crawler policy; allow answer-engine fetchers if GEO visibility is desired.", "GEO", "Low", "Very high"),
            ("Critical", "Force HTTP→HTTPS with 301/308 and add HSTS after verification.", "SEO", "Low", "High"),
            ("High", "Publish complete NAP and remove the empty streetAddress value.", "Local/AEO", "Low", "High"),
            ("High", "Create About/Team content with real people, credentials and sameAs profiles.", "GEO", "Medium", "High"),
            ("High", "Use unique OG metadata and a 1200×630 image per key template.", "SEO", "Low", "Medium"),
            ("Medium", "Expand Pricing/Areas and mark the construction process as HowTo.", "SEO/AEO", "Medium", "Medium"),
        ],
        "strengths": ["Unique FAQPage answers on every service", "Transparent price ranges", "Winter Haven neighborhood specificity", "Clean service URL and breadcrumb architecture"],
    },
    {
        "slug": "cleaningweekly-com",
        "name": "Cleaning Weekly",
        "domain": "cleaningweekly.com",
        "scores": (5, 4, 7),
        "pages": "All 52 sitemap URLs: homepage, booking, one pillar guide, Locations hub, 6 city pages, 36 city×service pages and 6 core service pages; robots.txt and both sitemap files inspected.",
        "summary": "Cleaning Weekly has a valuable local matrix and the portfolio's strongest single pricing guide, with TechArticle, FAQPage, Speakable and a comparison table. Critical launch blockers outweigh that foundation: every canonical and sitemap URL uses a www host that did not resolve during the audit; phone values are blank or fictional 555 numbers; sameAs contains REPLACE_PROFILE placeholders; and major AI crawlers are blocked. No About, Team or Contact page exists.",
        "findings": [
            ("SEO", "Canonical host", "All canonicals and sitemap URLs declare www, but the www host did not resolve while apex served content.", "Critical"),
            ("SEO", "NAP integrity", "Homepage phone is blank/'coming soon'; other pages and schema use (407) 555-0148.", "Critical"),
            ("SEO", "Architecture", "Six cities × six services have clean paths, unique titles and broad internal links.", "Good"),
            ("SEO", "Content depth", "Core service pages are about 260–275 words; Locations hub 185; Book page about 104.", "Needs Attention"),
            ("SEO", "Social metadata", "OG image is SVG, which is unreliable on major social platforms.", "Needs Attention"),
            ("GEO", "AI crawler access", "robots.txt blocks GPTBot, Google-Extended, ClaudeBot and related fetchers.", "Critical"),
            ("GEO", "Entity integrity", "sameAs contains REPLACE_PROFILE stubs and the entity has no real phone or street address.", "Critical"),
            ("GEO", "E-E-A-T", "About, Team and Contact routes are missing; most imagery is stock Unsplash.", "Missing"),
            ("AEO", "Pillar guide", "Orlando weekly pricing guide includes a direct answer, table, FAQPage, TechArticle and Speakable.", "Good"),
            ("AEO", "Coverage gaps", "Homepage and core-service FAQs lack FAQPage; no HowTo exists.", "Needs Attention"),
        ],
        "priorities": [
            ("Critical", "Make www resolve or change all canonical/sitemap/schema URLs to apex, then 301 the alternate host.", "SEO", "Medium", "Very high"),
            ("Critical", "Replace blank/555 phone values everywhere with the real business number.", "Local/GEO", "Low", "Very high"),
            ("Critical", "Remove every REPLACE_PROFILE sameAs entry until a real profile exists.", "GEO", "Low", "High"),
            ("High", "Allow AI fetchers if citations in AI answers are a business goal.", "GEO", "Low", "Very high"),
            ("High", "Add About and Contact pages with real entity and trust information.", "GEO", "Medium", "High"),
            ("High", "Add FAQPage to homepage/core services and HowTo to the three-step flow.", "AEO", "Low", "High"),
        ],
        "strengths": ["Six-city by six-service local matrix", "Excellent Orlando pricing guide", "TechArticle/FAQ/Speakable implementation", "Clear recurring-service value proposition"],
    },
    {
        "slug": "cleaningdavenport-com",
        "name": "Cleaning Davenport",
        "domain": "cleaningdavenport.com",
        "scores": (7, 4, 7),
        "pages": "35 content pages: homepage, About, Contact, FAQ, Services hub, 7 services, Areas hub, 7 area pages, Pricing, Booking, 4 frequency pages, Reviews, Gallery, Blog index and 6 posts; robots.txt and sitemap.xml inspected.",
        "summary": "Cleaning Davenport has a coherent services×areas×frequency architecture, good on-page hygiene, FAQ/Service/Breadcrumb schema, and a clear pay-after-cleaning proposition. Its GEO score is suppressed because robots.txt blocks major AI answer crawlers, the public phone still says 'Coming soon,' and LocalBusiness schema omits telephone, street address, geo, logo and sameAs. Blog posts are short and have almost no topical H2 structure; testimonials exist but are not marked up.",
        "findings": [
            ("SEO", "On-page basics", "Titles, descriptions, one H1, self-canonicals, index/follow, viewport and image alts are consistently implemented.", "Good"),
            ("SEO", "Duplicate title", "/services and /areas/davenport-fl share the same title.", "Needs Attention"),
            ("SEO", "Social metadata", "Core pages omit og:image; service/area/blog templates include it.", "Needs Attention"),
            ("SEO", "Editorial depth", "Six posts contain roughly 270–361 body words and almost no topical H2s.", "Needs Attention"),
            ("SEO", "Structured data", "LocalBusiness, Service, FAQPage, BlogPosting and BreadcrumbList are present.", "Good"),
            ("GEO", "AI crawler access", "robots.txt disallows GPTBot, ClaudeBot, Google-Extended and other AI fetchers.", "Critical"),
            ("GEO", "Local entity", "Phone is 'Coming soon'; schema lacks telephone, street/postal, geo, logo, image and sameAs.", "Critical"),
            ("GEO", "E-E-A-T", "About is process-focused with no named people, credentials or authentic first-party imagery.", "Needs Attention"),
            ("AEO", "FAQ system", "Homepage, FAQ, all service pages and all area pages expose real FAQPage answers.", "Good"),
            ("AEO", "Advanced markup", "Homepage's three-step process lacks HowTo; Speakable and comparison tables are absent.", "Missing"),
        ],
        "priorities": [
            ("Critical", "Allow AI crawlers if GEO citations are desired; retain training restrictions separately if needed.", "GEO", "Low", "Very high"),
            ("Critical", "Publish a real phone and complete the LocalBusiness entity fields.", "SEO/GEO/AEO", "Low–Medium", "Very high"),
            ("High", "Add policy-compliant Review/AggregateRating markup for genuine testimonials.", "SEO/GEO", "Medium", "High"),
            ("High", "Expand blog posts to pillar depth with question H2s and named human authors.", "SEO/GEO/AEO", "Medium–High", "High"),
            ("Medium", "Expand satellite area pages and use real Davenport job photography.", "SEO/GEO", "Medium", "Medium"),
            ("Quick Win", "Add default OG imagery, differentiate the duplicate title, and mark the process as HowTo.", "SEO/AEO", "Low", "Medium"),
        ],
        "strengths": ["Complete local-service content system", "Strong FAQPage coverage", "Clean technical SEO fundamentals", "Consistent pay-after-cleaning differentiation"],
    },
]


def score_color(score):
    return GREEN if score >= 8 else AMBER if score >= 5 else RED


def status(score):
    return "Strong" if score >= 8 else "On Track" if score >= 6 else "Needs Work"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=100, start=100, bottom=100, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=DARK, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True, "FFFFFF", 9)
        shade(table.rows[0].cells[i], NAVY)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False, DARK, 8)
            if row_index % 2:
                shade(cells[i], LIGHT_GRAY)
        if widths:
            for i, width in enumerate(widths):
                cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def configure_doc(doc, site):
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name, size, color in (
        ("Title", 30, "FFFFFF"),
        ("Heading 1", 22, NAVY),
        ("Heading 2", 16, BLUE),
        ("Heading 3", 12, DARK),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    header = section.header.paragraphs[0]
    header.text = f"{site['domain']}  ·  SEO / GEO / AEO Audit"
    header.style = styles["Normal"]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.text = f"Booking Broom portfolio audit · {DATE}"
    footer.style = styles["Normal"]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_docx_cover(doc, site):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, NAVY)
    margins(cell, 900, 400, 900, 400)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(site["domain"])
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(30)
    r.font.color.rgb = RGBColor(255, 255, 255)
    p = cell.add_paragraph("SEO / GEO / AEO Audit Report")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(18)
    p.runs[0].font.color.rgb = RGBColor.from_string("93C5FD")
    p = cell.add_paragraph(f"FULL AUDIT  ·  {DATE}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

    scores = doc.add_table(rows=1, cols=3)
    for i, (label, score) in enumerate(zip(("SEO", "GEO", "AEO"), site["scores"])):
        cell = scores.cell(0, i)
        shade(cell, score_color(score))
        margins(cell, 250, 100, 250, 100)
        set_cell_text(cell, f"{label}\n{score}/10\n{status(score)}", True, "FFFFFF", 15, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def build_docx(site):
    doc = Document()
    configure_doc(doc, site)
    add_docx_cover(doc, site)

    doc.add_heading("Executive Summary", level=1)
    summary_box = doc.add_table(rows=1, cols=1)
    shade(summary_box.cell(0, 0), LIGHT_BLUE)
    set_cell_text(summary_box.cell(0, 0), site["summary"], False, DARK, 10)
    doc.add_paragraph()

    rows = []
    for label, score in zip(("SEO", "GEO", "AEO"), site["scores"]):
        rows.append((label, f"{score}/10", status(score)))
    score_table = add_table(doc, ["Dimension", "Score", "Status"], rows)
    for row, score in zip(score_table.rows[1:], site["scores"]):
        shade(row.cells[1], score_color(score))
        for run in row.cells[1].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    doc.add_heading("Crawl Scope", level=1)
    doc.add_paragraph(site["pages"])
    doc.add_paragraph("Audit method: homepage + robots/sitemap discovery, followed by every meaningful public URL found in navigation, internal links, or sitemap. Legal, login, and confirmation pages were excluded. Confirmed 404 probes are reflected in findings where relevant.")

    doc.add_heading("Signal-by-Signal Findings", level=1)
    add_table(doc, ["Dimension", "Signal", "Evidence", "Status"], site["findings"], [0.7, 1.2, 4.2, 1.0])

    doc.add_heading("Priority Recommendations", level=1)
    priority_table = add_table(doc, ["Priority", "Action", "Dimension", "Effort", "Impact"], site["priorities"], [0.8, 3.8, 1.0, 0.8, 0.8])
    for row in priority_table.rows[1:]:
        label = row.cells[0].text
        fill = RED if label == "Critical" else "EA580C" if label == "High" else AMBER if label == "Medium" else GREEN
        shade(row.cells[0], fill)
        for run in row.cells[0].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    doc.add_heading("What's Working Well", level=1)
    good = doc.add_table(rows=1, cols=1)
    for item in site["strengths"]:
        cell = good.rows[-1].cells[0]
        shade(cell, LIGHT_GREEN)
        set_cell_text(cell, f"• {item}", False, DARK, 10)
        if item != site["strengths"][-1]:
            good.add_row()
    doc.add_paragraph()

    doc.add_heading("Measurement Limitations", level=1)
    doc.add_paragraph("This crawl evaluates delivered HTML, crawl directives, metadata, content and structured data. It does not measure backlink authority, Search Console index coverage, ranking positions, Google Business Profile completeness, conversion analytics, or field Core Web Vitals. Use PageSpeed Insights/CrUX for LCP, INP and CLS; Search Console for indexing and query performance; and a backlink platform for authority analysis.")

    doc.add_heading("Glossary", level=1)
    add_table(doc, ["Term", "Meaning"], [
        ("SEO", "Traditional search visibility: crawlability, metadata, content quality, internal links and structured data."),
        ("GEO", "Visibility and citation readiness in generative search and AI answer systems."),
        ("AEO", "Extractability for featured snippets, People Also Ask, assistants and voice queries."),
    ])

    path = OUT / f"seo-geo-aeo-full-audit-{site['slug']}-{DATE}.docx"
    doc.save(path)
    return path


def pdf_paragraph(text, style):
    return Paragraph(escape(str(text)).replace("\n", "<br/>"), style)


def pdf_table(headers, rows, widths, styles):
    data = [[pdf_paragraph(h, styles["TableHead"]) for h in headers]]
    for row in rows:
        data.append([pdf_paragraph(v, styles["TableCell"]) for v in row])
    table = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{NAVY}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor(f"#{MID_GRAY}")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(f"#{LIGHT_GRAY}")]),
    ]))
    return table


def build_pdf(site):
    path = OUT / f"seo-geo-aeo-full-audit-{site['slug']}-{DATE}.pdf"
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CoverDomain", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=27, leading=32, textColor=colors.white, alignment=TA_CENTER, spaceAfter=16))
    styles.add(ParagraphStyle(name="CoverSub", parent=styles["Normal"], fontName="Helvetica", fontSize=15, leading=20, textColor=colors.HexColor("#93C5FD"), alignment=TA_CENTER))
    styles.add(ParagraphStyle(name="H1x", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=20, leading=24, textColor=colors.HexColor(f"#{NAVY}"), spaceBefore=10, spaceAfter=9))
    styles.add(ParagraphStyle(name="Bodyx", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=14, textColor=colors.HexColor(f"#{DARK}")))
    styles.add(ParagraphStyle(name="TableHead", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=7.5, leading=9, textColor=colors.white))
    styles.add(ParagraphStyle(name="TableCell", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.2, leading=9, textColor=colors.HexColor(f"#{DARK}")))
    doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.6*inch, leftMargin=0.6*inch, topMargin=0.6*inch, bottomMargin=0.6*inch, title=f"{site['name']} SEO GEO AEO Audit")
    story = []
    cover = Table([[
        Paragraph(f"{escape(site['domain'])}<br/><font size='15' color='#93C5FD'>SEO / GEO / AEO Audit Report</font><br/><font size='10'>FULL AUDIT · {DATE}</font>", styles["CoverDomain"])
    ]], colWidths=[7.3*inch], rowHeights=[7.0*inch])
    cover.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{NAVY}")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("BOX", (0, 0), (-1, -1), 0, colors.HexColor(f"#{NAVY}")),
    ]))
    story.extend([Spacer(1, 0.7*inch), cover, PageBreak()])
    story.extend([Paragraph("Executive Summary", styles["H1x"]), pdf_paragraph(site["summary"], styles["Bodyx"]), Spacer(1, 10)])
    story.append(pdf_table(["Dimension", "Score", "Status"], [(label, f"{score}/10", status(score)) for label, score in zip(("SEO", "GEO", "AEO"), site["scores"])], [1.5*inch, 1.0*inch, 1.5*inch], styles))
    story.extend([Paragraph("Crawl Scope", styles["H1x"]), pdf_paragraph(site["pages"], styles["Bodyx"]), Spacer(1, 10)])
    story.extend([Paragraph("Signal-by-Signal Findings", styles["H1x"]), pdf_table(["Dimension", "Signal", "Evidence", "Status"], site["findings"], [0.65*inch, 1.15*inch, 4.3*inch, 1.0*inch], styles)])
    story.extend([Paragraph("Priority Recommendations", styles["H1x"]), pdf_table(["Priority", "Action", "Dimension", "Effort", "Impact"], site["priorities"], [0.75*inch, 3.65*inch, 1.0*inch, 0.8*inch, 0.8*inch], styles)])
    story.extend([Paragraph("What's Working Well", styles["H1x"]), pdf_table(["Verified strength"], [(x,) for x in site["strengths"]], [7.0*inch], styles)])
    story.extend([Paragraph("Measurement Limitations", styles["H1x"]), pdf_paragraph("This crawl evaluates delivered HTML, crawl directives, metadata, content and structured data. It does not measure backlinks, Search Console indexation, rankings, Google Business Profile performance, conversion analytics, or field Core Web Vitals.", styles["Bodyx"])])

    def footer(canvas, _doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#64748B"))
        canvas.drawString(0.6*inch, 0.35*inch, f"{site['domain']} · Full SEO/GEO/AEO audit")
        canvas.drawRightString(7.9*inch, 0.35*inch, f"Page {_doc.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    return path


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    generated = []
    for site in SITES:
        generated.extend([build_docx(site), build_pdf(site)])
    manifest = OUT / "README.md"
    lines = [
        f"# Booking Broom SEO / GEO / AEO Full Audits — {DATE}",
        "",
        "Seven websites were crawled from robots.txt and sitemap discovery through all meaningful public content.",
        "",
    ]
    for site in SITES:
        lines.append(f"- **{site['name']}** — SEO {site['scores'][0]}/10 · GEO {site['scores'][1]}/10 · AEO {site['scores'][2]}/10")
        lines.append(f"  - `seo-geo-aeo-full-audit-{site['slug']}-{DATE}.docx`")
        lines.append(f"  - `seo-geo-aeo-full-audit-{site['slug']}-{DATE}.pdf`")
    manifest.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(f"Generated {len(generated)} reports in {OUT}")
    for path in generated:
        print(path.name)


if __name__ == "__main__":
    main()
